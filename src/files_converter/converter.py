import os
import re
import shutil
import subprocess
import tarfile
import zipfile
import gettext

import docx
import ffmpeg
import PyPDF2
from pdf2docx import Converter
from striprtf.striprtf import rtf_to_text
from odf.opendocument import OpenDocumentText, load
from odf import teletype
from odf.text import P, Span
from odf.style import Style, TextProperties
from PIL import Image

_ = gettext.gettext


class FileConverter:
    def __init__(self):
        self.supported_formats = {
            "photos": ["jpg", "png", "gif", "bmp", "tiff", "webp"],
            "videos": ["mp4", "avi", "mov", "mkv", "webm"],
            "vectors": ["svg", "eps"],
            "audio": ["mp3", "wav", "ogg", "flac", "aac"],
            "documents": ["pdf", "docx", "txt", "rtf", "odt"],
            "archives": ["zip", "tar", "tar.gz", "tar.xz", "tar.bz2", "rar", "7z"],
            "ebooks": ["epub", "mobi", "azw3", "fb2", "txt", "rtf", "pdf"],
        }

    def get_file_type(self, file_path):
        extension = ".".join(os.path.basename(file_path).split(os.extsep)[1:])
        for file_type, formats in self.supported_formats.items():
            if extension in formats:
                return file_type
        return None

    def get_target_formats(self, file_type):
        return self.supported_formats.get(file_type, [])

    def get_current_format(self, file_path):
        return os.path.splitext(file_path)[1][1:].lower()

    def convert_file(self, input_path, output_path, target_format, progress_callback=None):
        self.progress_callback = progress_callback
        file_type = self.get_file_type(input_path)

        if file_type == "photos":
            self._convert_photo(input_path, output_path, target_format)
        elif file_type == "videos":
            self._convert_video(input_path, output_path, target_format)
        elif file_type == "vectors":
            self._convert_vector(input_path, output_path, target_format)
        elif file_type == "audio":
            self._convert_audio(input_path, output_path, target_format)
        elif file_type == "documents":
            self._convert_document(input_path, output_path, target_format)
        elif file_type == "archives":
            self._convert_archive(input_path, output_path, target_format)
        elif file_type == "ebooks":
            self._convert_ebook(input_path, output_path, target_format)
        else:
            raise ValueError(_(f"Unsupported file type: {file_type}"))

    def _convert_photo(self, input_path, output_path, target_format):
        with Image.open(input_path) as img:
            if target_format.lower() == "jpg":
                target_format = "JPEG"
                img = img.convert("RGB")
            elif target_format.lower() == "tiff":
                target_format = "TIFF"

            save_kwargs = {}
            if target_format.upper() == "TIFF":
                save_kwargs["compression"] = "tiff_deflate"

            img.save(output_path, format=target_format.upper(), **save_kwargs)

    def _convert_video(self, input_path, output_path, target_format):
        try:
            # Get video duration
            probe = ffmpeg.probe(input_path)
            duration = float(probe.get("format", {}).get("duration", 0))

            # If duration is not in format info, try to find it in streams
            if duration == 0:
                for stream in probe["streams"]:
                    if "duration" in stream:
                        duration = float(stream["duration"])
                        break

            # If still no duration found, use a default value or raise an exception
            if duration == 0:
                print(
                    "Warning: Couldn't determine video duration. Progress reporting may be inaccurate."
                )
                duration = 1  # Default to 1 second to avoid division by zero

            # Set up the conversion
            stream = ffmpeg.input(input_path)

            # Handle different target formats
            if target_format == "mkv":
                stream = ffmpeg.output(stream, output_path, vcodec="libx264", acodec="libvorbis")
            else:
                stream = ffmpeg.output(stream, output_path, format=target_format)

            stream = stream.overwrite_output()

            # Construct the ffmpeg command
            ffmpeg_cmd = ffmpeg.compile(stream)

            # Run the conversion with subprocess
            with subprocess.Popen(
                ffmpeg_cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                bufsize=1,
                universal_newlines=True,
            ) as process:
                # Monitor the conversion progress
                pattern = re.compile(r"time=(\d{2}):(\d{2}):(\d{2})\.\d{2}")
                for line in process.stdout:
                    match = pattern.search(line)
                    if match:
                        hours, minutes, seconds = map(int, match.groups())
                        time_processed = hours * 3600 + minutes * 60 + seconds
                        progress = time_processed / duration
                        if hasattr(self, "progress_callback") and callable(self.progress_callback):
                            self.progress_callback(progress)

                # Ensure the process is complete
                process.wait()

                if process.returncode != 0:
                    raise Exception(
                        _(f"ffmpeg process failed with return code: {process.returncode}")
                    )

        except ffmpeg.Error as e:
            print("stdout:", e.stdout.decode("utf8"))
            print("stderr:", e.stderr.decode("utf8"))
            raise
        except Exception as e:
            print(f"An error occurred: {str(e)}")
            raise

    def _convert_vector(self, input_path, output_path, target_format):
        if shutil.which("inkscape"):
            subprocess.run(["inkscape", input_path, f"--export-filename={output_path}"])
        else:
            raise RuntimeError(
                _("Inkscape is required for vector conversions but is not installed.")
            )

    def _convert_audio(self, input_path, output_path, target_format):
        if target_format == "aac":
            (
                ffmpeg.input(input_path)
                .output(output_path, acodec="aac", format="adts")
                .overwrite_output()
                .run()
            )
        else:
            (
                ffmpeg.input(input_path)
                .output(output_path, format=target_format)
                .overwrite_output()
                .run()
            )

    def _convert_document(self, input_path, output_path, target_format):
        if target_format == "pdf":
            if input_path.endswith(".docx"):
                doc = docx.Document(input_path)
                doc.save(output_path)
            elif input_path.endswith(".txt"):
                from reportlab.lib.pagesizes import letter
                from reportlab.pdfgen import canvas

                c = canvas.Canvas(output_path, pagesize=letter)
                with open(input_path, "r") as txt_file:
                    text = txt_file.read()
                    c.drawString(72, 800, text)
                c.save()
        elif target_format == "txt":
            if input_path.endswith(".pdf"):
                with open(input_path, "rb") as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    with open(output_path, "w") as txt_file:
                        for page in reader.pages:
                            txt_file.write(page.extract_text())
            elif input_path.endswith(".docx"):
                doc = docx.Document(input_path)
                with open(output_path, "w") as txt_file:
                    for para in doc.paragraphs:
                        txt_file.write(para.text + "\n")
        elif target_format == "docx":
            if input_path.endswith(".pdf"):
                cv = Converter(input_path)
                cv.convert(output_path)
                cv.close()
            elif input_path.endswith(".txt"):
                # Create a new DOCX document
                doc = docx.Document()

                # Read the TXT file and add its content to the DOCX
                with open(input_path, "r", encoding="utf-8") as txt_file:
                    for line in txt_file:
                        doc.add_paragraph(line.strip())

                # Save the DOCX file
                doc.save(output_path)
            elif input_path.endswith(".rtf"):
                # Read the RTF file
                with open(input_path, "r", encoding="utf-8") as rtf_file:
                    rtf_text = rtf_file.read()

                # Convert RTF to plain text
                plain_text = rtf_to_text(rtf_text)

                # Create a new DOCX document
                doc = docx.Document()

                # Add the content to the DOCX
                for paragraph in plain_text.split("\n"):
                    doc.add_paragraph(paragraph)

                # Save the DOCX file
                doc.save(output_path)
            elif input_path.endswith(".odt"):
                # Load the ODT document
                odt_doc = load(input_path)

                # Create a new DOCX document
                docx_doc = docx.Document()

                # Iterate through the ODT content
                for element in odt_doc.getElementsByType(P):
                    # Extract text from the paragraph
                    paragraph_text = teletype.extractText(element)

                    # Add the paragraph to the DOCX document
                    docx_doc.add_paragraph(paragraph_text)

                # Save the DOCX file
                docx_doc.save(output_path)
            else:
                raise ValueError(
                    _(f"Unsupported conversion to DOCX from {os.path.splitext(input_path)[1]}")
                )
        elif target_format == "rtf":
            if input_path.endswith(".pdf"):
                with open(input_path, "rb") as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text()

                rtf_content = "{\\rtf1\\ansi\n"
                rtf_content += text.replace("\n", "\\par\n")
                rtf_content += "}"

                with open(output_path, "w") as rtf_file:
                    rtf_file.write(rtf_content)
            elif input_path.endswith(".docx"):
                doc = docx.Document(input_path)
                rtf_content = "{\\rtf1\\ansi\n"
                for para in doc.paragraphs:
                    rtf_content += para.text.replace("\n", "\\par\n") + "\\par\n"
                rtf_content += "}"

                with open(output_path, "w", encoding="utf-8") as rtf_file:
                    rtf_file.write(rtf_content)
            elif input_path.endswith(".txt"):
                # Read the TXT file
                with open(input_path, "r", encoding="utf-8") as txt_file:
                    text_content = txt_file.read()

                # Create a simple RTF structure
                rtf_content = "{\\rtf1\\ansi\\deff0\n"
                rtf_content += "{\\fonttbl{\\f0 Arial;}}\n"
                rtf_content += "\\f0\\fs24\n"

                # Convert newlines to RTF paragraph breaks
                rtf_content += text_content.replace("\n", "\\par\n")

                rtf_content += "}"

                # Write the RTF content to the output file
                with open(output_path, "w", encoding="utf-8") as rtf_file:
                    rtf_file.write(rtf_content)
            elif input_path.endswith(".odt"):
                # Load the ODT document
                odt_doc = load(input_path)

                # Create RTF header
                rtf_content = "{\\rtf1\\ansi\\deff0\n"
                rtf_content += "{\\fonttbl{\\f0 Arial;}}\n"
                rtf_content += "\\f0\\fs24\n"

                # Iterate through the ODT content
                for element in odt_doc.getElementsByType(P):
                    # Extract text from the paragraph
                    paragraph_text = teletype.extractText(element)

                    # Add the paragraph to RTF content
                    rtf_content += paragraph_text.replace("\n", "\\par\n") + "\\par\n"

                # Close RTF document
                rtf_content += "}"

                # Write RTF content to file
                with open(output_path, "w", encoding="utf-8") as rtf_file:
                    rtf_file.write(rtf_content)
            else:
                raise ValueError(
                    _(f"Unsupported conversion to RTF from {os.path.splitext(input_path)[1]}")
                )
        elif target_format == "odt":
            if input_path.endswith(".pdf"):
                # Create a new ODT document
                doc = OpenDocumentText()

                # Read the PDF and extract text
                with open(input_path, "rb") as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    for page in reader.pages:
                        text = page.extract_text()
                        # Add each line as a paragraph in the ODT
                        for line in text.split("\n"):
                            p = P(text=line)
                            doc.text.addElement(p)

                # Save the ODT file
                doc.save(output_path)
            elif input_path.endswith(".docx"):
                # Create a new ODT document
                odt_doc = OpenDocumentText()

                # Read the DOCX file
                docx_doc = docx.Document(input_path)

                # Create some basic styles
                style_bold = Style(name="Bold", family="text")
                style_bold.addElement(TextProperties(fontweight="bold"))
                odt_doc.automaticstyles.addElement(style_bold)

                style_italic = Style(name="Italic", family="text")
                style_italic.addElement(TextProperties(fontstyle="italic"))
                odt_doc.automaticstyles.addElement(style_italic)

                # Convert paragraphs
                for paragraph in docx_doc.paragraphs:
                    odt_para = P()
                    for run in paragraph.runs:
                        span = Span(text=run.text)
                        if run.bold:
                            span.stylename = style_bold
                        if run.italic:
                            span.stylename = style_italic
                        odt_para.addElement(span)
                    odt_doc.text.addElement(odt_para)

                # Save the ODT file
                odt_doc.save(output_path)
            elif input_path.endswith(".txt"):
                # Create a new ODT document
                odt_doc = OpenDocumentText()

                # Create a default text style
                style_default = Style(name="Default", family="paragraph")
                style_default.addElement(TextProperties(fontname="Arial", fontsize="12pt"))
                odt_doc.automaticstyles.addElement(style_default)

                # Read the TXT file and add its content to the ODT
                with open(input_path, "r", encoding="utf-8") as txt_file:
                    for line in txt_file:
                        p = P(stylename=style_default)
                        p.addText(line.strip())
                        odt_doc.text.addElement(p)

                # Save the ODT file
                odt_doc.save(output_path)
            elif input_path.endswith(".rtf"):
                # Create a new ODT document
                odt_doc = OpenDocumentText()

                # Create a default text style
                style_default = Style(name="Default", family="paragraph")
                style_default.addElement(TextProperties(fontname="Arial", fontsize="12pt"))
                odt_doc.automaticstyles.addElement(style_default)

                # Read the RTF file
                with open(input_path, "r", encoding="utf-8") as rtf_file:
                    rtf_text = rtf_file.read()

                # Convert RTF to plain text
                plain_text = rtf_to_text(rtf_text)

                # Add the content to the ODT
                for paragraph in plain_text.split("\n"):
                    p = P(stylename=style_default)
                    p.addText(paragraph)
                    odt_doc.text.addElement(p)

                # Save the ODT file
                odt_doc.save(output_path)
            else:
                raise ValueError(
                    _(f"Unsupported conversion to ODT from {os.path.splitext(input_path)[1]}")
                )
        else:
            raise ValueError(
                _(
                    f"Unsupported document conversion: {os.path.splitext(input_path)[1]} to {target_format}"
                )
            )

    def _convert_archive(self, input_path, output_path, target_format):
        temp_dir = os.path.join(os.path.dirname(output_path), "temp_extract")
        os.makedirs(temp_dir, exist_ok=True)

        try:
            # Extract the input archive
            if input_path.endswith((".tar", ".tar.gz", ".tar.xz", ".tar.bz2")):
                with tarfile.open(input_path, "r:*") as tar:
                    tar.extractall(temp_dir)
            elif input_path.endswith(".zip"):
                with zipfile.ZipFile(input_path, "r") as zip_ref:
                    zip_ref.extractall(temp_dir)
            elif input_path.endswith(".rar"):
                subprocess.run(["unrar", "x", input_path, temp_dir], check=True)
            elif input_path.endswith(".7z"):
                subprocess.run(["7z", "x", input_path, f"-o{temp_dir}"], check=True)
            else:
                raise ValueError(
                    _(f"Unsupported input archive format: {os.path.splitext(input_path)[1]}")
                )

            # Create the output archive
            if target_format == "zip":
                base_name = os.path.basename(output_path).split(os.extsep)[0]
                shutil.make_archive(
                    os.path.join(os.path.dirname(output_path), base_name), "zip", temp_dir
                )
            elif target_format in ["tar", "tar.gz", "tar.xz", "tar.bz2"]:
                base_name = os.path.basename(output_path).split(os.extsep)[0]
                if target_format == "tar":
                    shutil.make_archive(
                        os.path.join(os.path.dirname(output_path), base_name), "tar", temp_dir
                    )
                elif target_format in ["tar.xz", "tar.bz2"]:
                    compression = "xz" if target_format == "tar.xz" else "bz2"
                    with tarfile.open(output_path, f"w:{compression}") as tar:
                        tar.add(temp_dir, arcname="")
                else:
                    shutil.make_archive(
                        os.path.join(os.path.dirname(output_path), base_name), "gztar", temp_dir
                    )
            elif target_format == "rar":
                subprocess.run(["rar", "a", output_path, temp_dir], check=True)
            elif target_format == "7z":
                subprocess.run(["7z", "a", output_path, temp_dir], check=True)
            else:
                raise ValueError(_(f"Unsupported output archive format: {target_format}"))

        except subprocess.CalledProcessError as e:
            raise RuntimeError(_(f"Error during archive conversion: {str(e)}"))
        except Exception as e:
            raise ValueError(_(f"Error during archive conversion: {str(e)}"))
        finally:
            # Clean up temporary directory
            shutil.rmtree(temp_dir, ignore_errors=True)

    def _convert_ebook(self, input_path, output_path, target_format):
        try:
            subprocess.run(["ebook-convert", input_path, output_path], check=True)
        except subprocess.CalledProcessError:
            raise RuntimeError(_(f"Failed to convert ebook from {input_path} to {output_path}"))
